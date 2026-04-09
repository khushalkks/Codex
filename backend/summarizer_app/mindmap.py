# import io
# import requests
# import json
# import re
# import PyPDF2
# import docx
# from summarizer_app.rag_service import create_embeddings, retrieve

# # ── Config ────────────────────────────────────────────────────────────────────

# OLLAMA_URL = "http://localhost:11434/api/generate"
# MODEL = "mistral"


# # ── File Text Extraction ──────────────────────────────────────────────────────

# def extract_text_from_file(file) -> str:
#     """
#     Extract plain text from:
#       - io.BytesIO object  (with .filename attribute set by routes.py)
#       - file path string   (used by summarizer)
#       - FastAPI UploadFile (with .filename attribute)
#     """

#     # ── Case 1: file path string (e.g. from summarizer) ──────────────────────
#     if isinstance(file, str):
#         filename = file.lower()
#         if filename.endswith(".pdf"):
#             with open(file, "rb") as f:
#                 reader = PyPDF2.PdfReader(f)
#                 return "\n".join(page.extract_text() or "" for page in reader.pages)
#         elif filename.endswith(".docx"):
#             doc = docx.Document(file)
#             return "\n".join(para.text for para in doc.paragraphs)
#         elif filename.endswith(".txt"):
#             with open(file, "r", encoding="utf-8", errors="ignore") as f:
#                 return f.read()
#         else:
#             raise ValueError(f"Unsupported file type: {file}")

#     # ── Case 2: BytesIO or UploadFile (has .filename attribute) ──────────────
#     filename = getattr(file, "filename", "").lower()

#     # BytesIO se pehle pointer reset karo
#     if isinstance(file, io.BytesIO):
#         file.seek(0)

#     if filename.endswith(".pdf"):
#         reader = PyPDF2.PdfReader(file)
#         return "\n".join(page.extract_text() or "" for page in reader.pages)

#     elif filename.endswith(".docx"):
#         doc = docx.Document(file)
#         return "\n".join(para.text for para in doc.paragraphs)

#     elif filename.endswith(".txt"):
#         raw = file.read()
#         return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else raw

#     else:
#         raise ValueError(f"Unsupported file type: {filename or 'unknown'}")


# # ── Keyword Extraction via Ollama ─────────────────────────────────────────────

# def extract_keywords(text: str) -> list[str]:
#     """Ask the LLM to pull the top keywords from the document."""
#     snippet = text[:4000]

#     prompt = f"""
# Extract the 5 most important keywords or key concepts from the following text.

# Return ONLY a JSON array of strings. No explanation, no markdown.

# Example output: ["Machine Learning", "Neural Networks", "Training Data", "Overfitting", "Gradient Descent"]

# Text:
# {snippet}
# """

#     payload = {
#         "model": MODEL,
#         "prompt": prompt,
#         "stream": False,
#         "options": {"temperature": 0.3}
#     }

#     resp = requests.post(OLLAMA_URL, json=payload, timeout=60)
#     raw = resp.json().get("response", "[]")

#     # Strip markdown fences if present
#     raw = re.sub(r"```[a-z]*", "", raw).strip().strip("`").strip()

#     try:
#         keywords = json.loads(raw)
#         if isinstance(keywords, list):
#             return [str(k) for k in keywords[:5]]
#     except Exception:
#         pass

#     # Fallback: grab quoted strings
#     return re.findall(r'"([^"]+)"', raw)[:5]


# # ── Mindmap Tree Generation via Ollama ────────────────────────────────────────

# def generate_tree(topic: str) -> dict:
#     """Ask the LLM to create a nested knowledge tree for a topic."""
#     prompt = f"""
# Create a structured knowledge tree for the topic.

# Return ONLY JSON — no explanation, no markdown fences.

# Format:
# {{
#   "name": "{topic}",
#   "children": [
#     {{
#       "name": "Concept",
#       "children": [
#         {{"name": "Subtopic"}},
#         {{"name": "Subtopic"}}
#       ]
#     }}
#   ]
# }}

# Topic: {topic}

# Rules:
# - Create exactly 3 main branches
# - Each branch must have exactly 3 subtopics
# - Return JSON only
# """

#     payload = {
#         "model": MODEL,
#         "prompt": prompt,
#         "stream": False,
#         "options": {"temperature": 0.7}
#     }

#     resp = requests.post(OLLAMA_URL, json=payload, timeout=90)
#     raw = resp.json().get("response", "")
#     raw = re.sub(r"```[a-z]*", "", raw).strip().strip("`").strip()

#     return json.loads(raw)


# # ── JSON → Mermaid Converter ──────────────────────────────────────────────────

# def json_to_mermaid(tree: dict) -> str:
#     lines = ["mindmap"]

#     def walk(node, level=1):
#         indent = "  " * level
#         name = node.get("name", "?")
#         if level == 1:
#             lines.append(f"{indent}root(({name}))")
#         elif level == 2:
#             lines.append(f"{indent}[{name}]")
#         else:
#             lines.append(f"{indent}{name}")
#         for child in node.get("children", []):
#             walk(child, level + 1)

#     walk(tree)
#     return "\n".join(lines)


# # ── Main pipeline function (called from routes.py) ────────────────────────────

# def generate_mindmap(file) -> dict:
#     """Full pipeline: file → keywords → mermaid mindmap."""
#     text = extract_text_from_file(file)
#     keywords = extract_keywords(text)
#     primary = keywords[0] if keywords else "Document"
#     tree = generate_tree(primary)
#     mermaid = json_to_mermaid(tree)

#     return {
#         "keywords": keywords,
#         "primary_topic": primary,
#         "mermaid": mermaid,
#         "tree": tree
#     }

import io
import requests
import json
import re
import PyPDF2
import docx
from summarizer_app.rag_service import create_embeddings, retrieve

# ── CONFIG ─────────────────────────────────────────────

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL = "mistral"

# ── FILE TEXT EXTRACTION ─────────────────────────────

def extract_text_from_file(file) -> str:

    if isinstance(file, str):
        filename = file.lower()

        if filename.endswith(".pdf"):
            with open(file, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.endswith(".docx"):
            doc = docx.Document(file)
            return "\n".join(p.text for p in doc.paragraphs)

        elif filename.endswith(".txt"):
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        else:
            raise ValueError("Unsupported file type")

    filename = getattr(file, "filename", "").lower()

    if isinstance(file, io.BytesIO):
        file.seek(0)

    if filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif filename.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join(p.text for p in doc.paragraphs)

    elif filename.endswith(".txt"):
        raw = file.read()
        return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else raw

    else:
        raise ValueError("Unsupported file type")


# ── 🔥 FIXED KEYWORD EXTRACTION ───────────────────────

def extract_keywords(text: str) -> list[str]:

    snippet = text[:4000]

    prompt = f"""
Extract exactly 5 important keywords from the text.

Return ONLY a valid JSON array.
Do NOT add explanation.
Do NOT add text before or after.

Example:
["AI", "Machine Learning", "Neural Networks"]

Text:
{snippet}
"""

    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": 0.3}
    }

    resp = requests.post(OLLAMA_URL, json=payload, timeout=60)
    raw = resp.json().get("response", "")

    print("RAW KEYWORDS:", raw)  # DEBUG

    raw = re.sub(r"```[a-z]*", "", raw).strip().strip("`")

    # extract JSON safely
    match = re.search(r"\[.*?\]", raw, re.DOTALL)

    if match:
        try:
            return json.loads(match.group())
        except:
            pass

    # fallback (strong)
    keywords = re.findall(r'\b[A-Za-z][A-Za-z ]{2,}\b', raw)

    return keywords[:5] if keywords else ["General Topic"]


# ── TREE GENERATION (RAG ENABLED) ─────────────────────

def generate_tree(context: str) -> dict:

    prompt = f"""
Use this context to generate a structured mind map.

Context:
{context}

Rules:
- 3 main branches
- Each branch 3 subtopics
- Keep names short
- Return ONLY valid JSON
"""

    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": 0.6}
    }

    resp = requests.post(OLLAMA_URL, json=payload, timeout=90)
    raw = resp.json().get("response", "")

    print("RAW TREE:", raw)  # DEBUG

    raw = re.sub(r"```[a-z]*", "", raw).strip().strip("`")

    match = re.search(r"\{.*\}", raw, re.DOTALL)

    if match:
        try:
            return json.loads(match.group())
        except:
            pass

    return {
        "name": "Error",
        "children": [{"name": "Parsing Failed"}]
    }


# ── MERMAID CONVERTER ───────────────────────────────

def json_to_mermaid(tree: dict) -> str:

    lines = ["mindmap"]

    def walk(node, level=1):
        indent = "  " * level
        name = node.get("name", "?")

        if level == 1:
            lines.append(f"{indent}root(({name}))")
        elif level == 2:
            lines.append(f"{indent}[{name}]")
        else:
            lines.append(f"{indent}{name}")

        for child in node.get("children", []):
            walk(child, level + 1)

    walk(tree)
    return "\n".join(lines)


# ── 🔥 MAIN FUNCTION (RAG PIPELINE) ───────────────────

def generate_mindmap(file) -> dict:

    # 1️⃣ Extract text
    text = extract_text_from_file(file)

    if not text.strip():
        return {"error": "Empty file"}

    # 2️⃣ Create embeddings (RAG)
    create_embeddings(text)

    # 3️⃣ Extract keywords
    keywords = extract_keywords(text)

    primary = keywords[0] if keywords else "Document"

    # 4️⃣ Retrieve relevant chunks
    context_chunks = retrieve(primary)

    if not context_chunks:
        context_chunks = [text[:1000]]

    context = " ".join(context_chunks)

    # 5️⃣ Generate tree
    tree = generate_tree(context)

    # 6️⃣ Convert to mermaid
    mermaid = json_to_mermaid(tree)

    return {
        "keywords": keywords,
        "primary_topic": primary,
        "mermaid": mermaid,
        "tree": tree
    }